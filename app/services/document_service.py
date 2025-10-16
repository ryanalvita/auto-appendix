from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from fastapi import UploadFile
import tempfile
import os
import subprocess


class DocumentService:
    # Paper sizes in centimeters (width, height)
    PAPER_SIZES = {
        "A4": (21.0, 29.7),
        "Letter": (21.6, 27.9),
        "Legal": (21.6, 35.6),
    }

    @staticmethod
    def create_appendix(
        files: list[UploadFile],
        image_width: float = 15,
        paper_size: str = "A4",
        output_format: str = "docx",
        caption_position: str = "bottom",
    ) -> str:
        """
        Create an appendix document with images.

        Args:
            files: List of uploaded image files
            image_width: Width of images in centimeters (default: 15)
            paper_size: Paper size ('A4', 'Letter', or 'Legal',
                default: 'A4')
            output_format: Output format ('docx' or 'pdf',
                default: 'docx')
            caption_position: Position of captions ('top' or 'bottom',
                default: 'bottom')

        Returns:
            Path to the generated document
        """
        doc = Document()

        # Set paper size
        if paper_size in DocumentService.PAPER_SIZES:
            width_cm, height_cm = DocumentService.PAPER_SIZES[paper_size]
            section = doc.sections[0]
            section.page_width = Cm(width_cm)
            section.page_height = Cm(height_cm)

        # Add 'Appendix' with Title style
        title = doc.add_paragraph("Appendix")
        title.style = "Title"
        doc.add_page_break()

        for idx, file in enumerate(files, start=1):
            filename_without_ext = Path(file.filename).stem

            # Caption above image
            if caption_position == "top":
                caption = doc.add_paragraph(f"Figure {idx}: {filename_without_ext}")
                caption.style = "Caption"

            # Add image (centered)
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()

            with tempfile.NamedTemporaryFile(
                delete=False, suffix=Path(file.filename).suffix
            ) as tmp:
                tmp.write(file.file.read())
                run.add_picture(tmp.name, width=Cm(image_width))
                os.unlink(tmp.name)

            # Caption below image
            if caption_position == "bottom":
                caption = doc.add_paragraph(f"Figure {idx}: {filename_without_ext}")
                caption.style = "Caption"

            # Page break (except for last image)
            if idx < len(files):
                doc.add_page_break()

        # Save as DOCX
        docx_path = tempfile.mktemp(suffix=".docx")
        doc.save(docx_path)

        # Convert to PDF if requested
        if output_format == "pdf":
            try:
                # Create a temporary directory for PDF conversion
                temp_dir = tempfile.mkdtemp()

                # Try using LibreOffice for conversion
                subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        temp_dir,
                        docx_path,
                    ],
                    check=True,
                    capture_output=True,
                    timeout=30,
                )

                # LibreOffice creates file with same name but .pdf ext
                converted_pdf = os.path.join(temp_dir, Path(docx_path).stem + ".pdf")

                if os.path.exists(converted_pdf):
                    # Move the PDF to a final location
                    pdf_path = tempfile.mktemp(suffix=".pdf")
                    os.rename(converted_pdf, pdf_path)
                    os.unlink(docx_path)
                    os.rmdir(temp_dir)
                    return pdf_path
                else:
                    # Clean up and raise error
                    os.unlink(docx_path)
                    os.rmdir(temp_dir)
                    raise Exception("PDF conversion failed. Output file not found.")
            except (
                subprocess.CalledProcessError,
                FileNotFoundError,
                subprocess.TimeoutExpired,
            ) as e:
                # If LibreOffice fails, clean up and raise error
                if os.path.exists(docx_path):
                    os.unlink(docx_path)
                if os.path.exists(temp_dir):
                    os.rmdir(temp_dir)
                raise Exception(
                    f"PDF conversion failed: {str(e)}. LibreOffice is "
                    "required for PDF export. Please install it or use "
                    "DOCX format."
                )

        return docx_path
